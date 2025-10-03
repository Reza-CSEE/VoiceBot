import os
import time
import telebot
import telegram
from telegram import Update, Bot
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, ConversationHandler
import stripe
import pymysql
import mysql.connector
import logging
import sounddevice as sd
import soundfile as sf
import speech_recognition as sr
import openai
from gtts import gTTS
import docx
from docx2pdf import convert
from ratelimit import limits, sleep_and_retry
import logging
import requests
from pydub import AudioSegment
import random
from datetime import datetime, timedelta


# Specify the paths to the ffmpeg and ffprobe executables
ffmpeg_path = "********/ffmpeg"
ffprobe_path = "********/ffprobe"
os.environ["PATH"] += os.pathsep + os.path.dirname(ffmpeg_path) + os.pathsep + os.path.dirname(ffprobe_path)

global my_voice_path 
my_voice_path = "********"



# Configure the logging
logging.basicConfig(level=logging.DEBUG, filename='app.log', filemode='w',
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Create a logger
logger = logging.getLogger()


# Set up the audio recording parameters
duration = 10  # seconds.  Agar in ro kheili ziyad bezaram va tedad jomle bala bere ehtemal javab eshtebah hast.
sample_rate = 44100
channels = 1

# Set up the OpenAI API key and model ID
openai.api_key = "********"
model_id = "********"

# Set your Stripe API key here
stripe.api_key = "********"


BOT_TOKEN = "********"

# Set up the Telegram bot

bot = Bot(token=BOT_TOKEN)

# @zebra_test_431  is for me.
chat = bot.get_chat('@********')
#chat = bot.get_chat('@********')
chat_id = chat.id


# Set up the Telegram bot
bot = telebot.TeleBot(BOT_TOKEN)

# Set up database connection
# connection = pymysql.connect(
#     host='********.rds.amazonaws.com',
#     port= ********,
#     user='********',
#     password='********',
#     database='********',
#     charset='********',
#     cursorclass=pymysql.cursors.DictCursor
# )


# # Create cursor
# cur = connection.cursor()

# #cur.execute("DELETE FROM  users")   # Delete all rows from the "users" table.     
# #cur.execute("DROP TABLE users")  # Delete the "users" table
# cur.execute("CREATE TABLE IF NOT EXISTS users (\
#     id INT AUTO_INCREMENT PRIMARY KEY, \
#     email VARCHAR(255), \
#     password VARCHAR(255), \
#     remaining_duration INT, \
#     customer_id  VARCHAR(255), \
#     payment_method_id  VARCHAR(255), \
#     plan_id VARCHAR(255) \
# )")

# cur.close()  # Close database cursor
# connection.commit()
# connection.close()  # Close database connection


# Create a dictionary to store the conversation history
conversation = {}
user_conversation = {}
bot_conversation = {}
bot_conversation_q ={}

# Define the rate-limiting rules
RATE_LIMIT = 6  # Limit 6 requests per minute

# Define a rate-limited function using the "limits" decorator
@sleep_and_retry
@limits(calls=RATE_LIMIT, period=60)
def handle_message(update: Update, context):
    global my_voice_path 
    user_id = update.effective_user.id
    chat_id = update.message.chat_id
    chat_member = bot.get_chat_member(chat_id=chat_id, user_id=user_id)
    if chat_member.status == 'member':
        try:    
            message = update.message
            # connection = pymysql.connect(
            # host='********.rds.amazonaws.com',
            # port= ********,
            # user='********',
            # password='********',
            # database='********',
            # charset='********',
            # cursorclass=pymysql.cursors.DictCursor)

            # Check if the user has an existing conversation
            if message.chat.id not in conversation:
                conversation[message.chat.id] = []
                user_conversation[message.chat.id] = []
                bot_conversation[message.chat.id] =[]
                bot_conversation_q[message.chat.id] =[]
            
            
            # Define a function for getting a response
            def get_response(prompt):
                global my_voice_path 
                if prompt == 'Please ask me a question.':
                    prompts1 = [
                    "What's something exciting or interesting that's happened to you recently?",
                    "What's your favorite way to spend your free time or relax?",
                    "If you could travel anywhere in the world, where would you go and why?",
                    "What's a book, movie, or TV show that you've been enjoying lately?",
                    "What's a skill or hobby you've always wanted to learn but haven't had the chance to yet?",
                    "Tell me about a memorable or funny experience you've had.",
                    "What's something you're passionate about or deeply interested in?",
                    "If you could have dinner with any historical figure, who would it be and what would you ask them?",
                    "What's a goal or aspiration you have for the future?",
                    "What's the most fascinating fact or piece of trivia you've come across recently?",
                    "What's something that always brings a smile to your face?",
                    "If you could have any superpower, what would it be and how would you use it?",
                    "What's a goal you've achieved recently that you're proud of?",
                    "What's the most memorable trip or vacation you've ever taken?",
                    "What's a skill or talent you have that not many people know about?",
                    "What's your favorite way to unwind after a long day?",
                    "If you could invite three people, dead or alive, to a dinner party, who would they be?",
                    "What's something new or interesting you've learned recently?",
                    "If you could live in any era of history, which one would you choose and why?",
                    "What's a quote or mantra that inspires you or resonates with you?",
                    "What's a skill or talent you admire in others and wish you had?",
                    "What's the best piece of advice you've ever received?",
                    "If you could have a conversation with anyone, living or deceased, who would it be and why?",
                    "What's something you've always wanted to try but haven't had the opportunity to do yet?",
                    "If you could change one thing about the world, what would it be?",
                    "What's your favorite way to challenge yourself and step out of your comfort zone?",
                    "What's a random act of kindness someone has done for you that you'll never forget?",
                    "What's a song or type of music that instantly puts you in a good mood?",
                    "If you could master one new language instantly, which one would you choose and why?",
                    "What's something you're currently working on or striving towards in your personal or professional life?",
                    "What's a memorable adventure or outdoor activity you've experienced?",
                    "If you could have a conversation with your younger self, what advice would you give?",
                    "What's a skill or hobby you've always wanted to learn but haven't had the time for?",
                    "What's something that never fails to make you laugh or smile?",
                    "If you could witness any historical event, which one would you choose and why?",
                    "What's a book, movie, or TV show that has had a significant impact on you and why?",
                    "What's the most interesting or unusual job you've ever had or heard of?",
                    "If you could have a home anywhere in the world, where would it be and what would it look like?",
                    "What's a cause or social issue that you're passionate about?",
                    "If you could spend a day with a fictional character, who would you choose and what would you do together?"
                    ]

                    # Randomly select a prompt
                    selected_prompt1 = random.choice(prompts1)
                    response_text2 = selected_prompt1 
                    print(response_text2)



                else:
                    response = openai.Completion.create(
                    engine=model_id,
                    prompt=prompt,
                    max_tokens=80,
                    n=1,
                    stop=None,
                    temperature=0.5,
                    )
                    # Print response to debug
                    #print(response)
                    # Get the response text from the OpenAI API
                    #response_text2 = "" 
                    response_text2 = response.choices[0].text.strip()
                    lines = response_text2.split("\n")
                    if len(lines) >= 2:
                        if lines[0].strip() == "" or lines[1].strip() == "":
                            response_text2 = "".join(lines[2:])
                    # Print response_text to debug
                    print(response_text2)
                # Convert the response text to speech using gTTS
                
                # You can change lang="en" for different languages 
                tts2 = gTTS(text=response_text2, lang="en", tld="com", slow=False)
                tts2.save(os.path.join(my_voice_path, "custom_voice2.mp3"))
                # Send the response text as an audio message
                with open(os.path.join(my_voice_path, "custom_voice2.mp3"), 'rb') as voice:
                    bot.send_audio(message.chat.id, voice)
                
                # Delete the file after sending
                os.remove(os.path.join(my_voice_path, "custom_voice2.mp3"))
                conversation[message.chat.id].append("Elisa: " + response_text2)
                
            conversation_history = "\n".join(conversation[message.chat.id])
            
            ########################################################################################################
            
            if message.text == "/help":
                logger.info("Received /help command")
                # Display help options
                bot.reply_to(message, "Here are the help options:\n"
                            "1. Record your voice to start the conversation\n"
                            "2. /ask - Ask me a question\n"
                            "3. /grammar - Correct your conversation grammatically\n"
                            "4. /score - Score the conversation like an IELTS marker\n"   
                            "5. /analyze - analyze the conversation like an IELTS marker\n"   
                            "6. /words - Give me the advanced words in this conversation\n"
                            "7. /text - Show me the text of our conversation\n"
                            "8. /export - Export word and PDF files of conversation history\n"
                            "9. /sign_up\n"
                            "10. /sign_in\n"
                            "11. /subscribe - You can pay a subscription\n"
                            "12. /links - Show useful links\n"   
                            "13. /help - Display help options\n"
                            "14. /support - Show support information\n"
                            "15. /privacy - Display privacy policy "
                            )
                
            elif message.text == "/ielts":
                    context.user_data["ielts"] = True
                    context.user_data["flag"] = False
                    user_id = update.effective_user.id
                    chat_id = update.effective_chat.id             
                    function1(user_id, chat_id)
                    


            

            elif message.text == "/links":
                logger.info("Received /links command")
                bot.reply_to(message, "Useful links:\n\n"
                                    "The most used words in the English language:\n"
                                    "https://gonaturalenglish.com/1000-most-common-words-in-the-english-language/\n"
                                    "https://www.vocabulary.com/lists/154147\n\n"
                                    "List of most commonly used English sentences:\n"
                                    "https://www.mondly.com/english-phrases-expressions\n"
                                    "https://www.krisamerikos.com/blog/english-sentences-used-in-daily-life\n"
                                    "https://www.fluentu.com/blog/english/basic-english-phrases/")
            elif message.text == "/support":
                logger.info("Received /support command")
                bot.reply_to(message, "Please contact this channel for support:\n"
                                    "@ZebraZebraAI")
            elif message.text == "/privacy":
                logger.info("Received /privacy command")
                bot.reply_to(message, "Please read our privacy policy file:")
                with open("privacy policy.pdf", "rb") as file:
                    bot.send_document(chat_id=message.chat.id, document=file)

            ###################################################################################################################################
            elif message.text == "/sign_in":
                logger.info("Received /sign_in command")
                context.user_data['email_prompted'] = True  # set a flag to indicate that we're prompting for email
                bot.reply_to(message, "Please enter your email.") 
            elif 'email_prompted' in context.user_data:  # check if we've prompted for email
                email = message.text  # get email from user input
                context.user_data['email'] = email
                print(f"Email entered for signing in: {email}")
                del context.user_data['email_prompted']  # remove the flag
                # send a message to the user with their email
                user_id = message.chat_id
                bot.send_message(chat_id=user_id, text=f"Your email is {email}")
                cur = connection.cursor()
                X = cur.execute("SELECT * FROM users WHERE email=%s", (email,))
                print(cur.mogrify("SELECT * FROM users WHERE email=%s", (email,)))  # This line will print the actual SQL query that is being executed
                print("X is:",X)
                result = cur.fetchone()
                cur.close() 
                print("result is:",result)
                #if user exists, prompt for password
                if result:
                    context.user_data['user_id'] = result['id']
                    context.user_data['password_match_requested'] = True
                    update.message.reply_text("Please enter your password.")           
                else:
                    update.message.reply_text("Sorry, this email is not registered. Please /sign_up first.")
                return None
            elif 'password_match_requested' in context.user_data: 
                password_match = message.text  # get password from user input
                context.user_data['password_match_requested'] = password_match
                print(f"Password entered for fetching data: {password_match}")
                del context.user_data['password_match_requested']  # remove the flag
                # Check if the password matches
                password = password_match  
                email = context.user_data['email']    
                cur = connection.cursor()
                cur.execute("SELECT * FROM users WHERE email=%s AND password=%s", (email, password))
                result = cur.fetchone()
                cur.close() 
                user_id = context.user_data['user_id']
                
                if result:
                    update.message.reply_text("You have successfully signed in. Please enter /x to start your conversation")
                    return None
                else:
                    update.message.reply_text("Sorry, the password you entered is incorrect. Please try again.")
                    return None

            ###################################################################################################################################
            elif message.text == "/sign_up":
                logger.info("Received /sign_up command")
                context.user_data['email_requested'] = True  # set a flag to indicate that we're prompting for email
                bot.reply_to(message, "Please enter your email.") 
            elif 'email_requested' in context.user_data:  # check if we've prompted for email
                email_signup = message.text  # get email from user input
                context.user_data['email_requested'] = email_signup
                print(f"Email entered for signing up: {email_signup}")
                del context.user_data['email_requested']  # remove the flag
                context.user_data['email'] = email_signup
                # Check if the user already exists in the database
                email = email_signup 
                cur = connection.cursor()
                cur.execute("SELECT * FROM users WHERE email=%s", (email,))
                result = cur.fetchone()
                cur.close() 
                if result:
                    update.message.reply_text("You have already signed up.\nPlease /sign_in to continue.")
                    return None
                else:
                    context.user_data['password_requested'] = True  # set a flag to indicate that we're prompting for password
                    bot.reply_to(message, "Please create your password.") 
            ###################################################################################################################################
            elif 'password_requested' in context.user_data:  # check if we've prompted for password
                password = message.text  # get password from user input
                context.user_data['password_requested'] = password
                print(f"Password entered for signing up: {password}")
                del context.user_data['password_requested']  # remove the flag

                # Insert the user's information into the database
                email = context.user_data['email']
                cur = connection.cursor()
                cur.execute("INSERT INTO users (email, password, remaining_duration) VALUES (%s, %s, %s)", (email, password, 15))
                cur.execute("SELECT * FROM users WHERE email=%s", (email,))
                result = cur.fetchone()
                print(result)
                cur.close() 
                context.user_data['user_id'] = result['id'] 

                connection.commit()
                update.message.reply_text("You signed up successfully.\nPlease /sign_in to continue.")
                return None
            ###################################################################################################################################
            elif message.text == "/subscribe":
                logger.info("Received /subscribe command")
                #customer_id = "********"
                #payment_method_id = "********"
                #plan_id = "********"
                try:

                    if 'user_id' in context.user_data and context.user_data['user_id'] is not None:
                        user_id = context.user_data['user_id'] 
                        print("user_id = ", user_id)
                    else:
                        update.message.reply_text("Please choose the /subscribe menu after /sign_in.")
                        return None



                    if 'payment_method_id' in context.user_data and context.user_data['payment_method_id'] is not None:
                        cur = connection.cursor()
                        cur.execute("SELECT payment_method_id FROM users WHERE id = %s", (user_id,))
                        result = cur.fetchone()
                        print(result)
                        cur.close() 
                        payment_method_id = result['payment_method_id']
                        print("payment_method_id existed:",payment_method_id)
                    else:
                        payment_method = stripe.PaymentMethod.create(
                        type='card',
                        card={
                        'number': '********',
                        'exp_month': ********,
                        'exp_year': ********,
                        'cvc': '********',
                        },)
                        payment_method_id = payment_method.id
                        context.user_data['payment_method_id'] = payment_method_id
                        cur = connection.cursor()
                        cur.execute("UPDATE users SET payment_method_id = %s WHERE id = %s", (payment_method_id, user_id))
                        connection.commit()
                        cur.close() 
                        print("payment_method_id produced:",payment_method_id)
                        
                    
                    if 'plan_id' in context.user_data and context.user_data['plan_id'] is not None:
                        cur = connection.cursor()
                        cur.execute("SELECT plan_id FROM users WHERE id = %s", (user_id,))
                        result = cur.fetchone()
                        print(result)
                        cur.close() 
                        plan_id = result['plan_id']
                        print("plan_id existed:",plan_id)
                    else:
                        plan = stripe.Plan.create(
                        amount=2000, # It is Cents. so 1000 cents is equal to 10 cad.
                        currency='cad',
                        interval='month',
                        product={
                            'name': 'Monthly Subscription',
                            },
                        nickname='Monthly Plan',
                        )
                        plan_id = plan.id
                        context.user_data['plan_id'] = plan_id
                        cur = connection.cursor()
                        cur.execute("UPDATE users SET plan_id = %s WHERE id = %s", (plan_id, user_id))
                        connection.commit()
                        cur.close() 
                        print("plan_id produced:",plan_id)
                    



                    # Create a new customer in Stripe and attach the payment method
                    customer = stripe.Customer.create(payment_method=payment_method_id)
                    customer_id = customer.id
                    cur = connection.cursor()
                    cur.execute("UPDATE users SET customer_id = %s WHERE id = %s", (customer_id, user_id))
                    connection.commit()
                    cur.close() 
                    print("customer_id produced:",customer_id)



                    payment_intent = stripe.PaymentIntent.create(
                    amount=2000,  # It is Cents. so 1000 cents is equal to 10 cad.
                    currency='cad',
                    payment_method='********',
                    customer='********',
                    description='Example PaymentIntent',
                    confirm=True,
                    )

                    subscription = stripe.Subscription.create(
                    customer = customer.id,
                    items=[{"plan": plan_id}],
                    default_payment_method=payment_method_id,
                    expand=["latest_invoice.payment_intent"],
                    )
                
                    checkout_session = stripe.checkout.Session.create(
                    customer=customer.id,
                    payment_method_types=["card"],
                    line_items=[
                        {
                            "price": "********",
                            "quantity": 1
                        }
                    ],
                    mode="subscription",
                    success_url="********",
                    #success_url="********",
                    #cancel_url="********",
                )
                    checkout_url = checkout_session.url
                    bot.send_message(chat_id=update.effective_chat.id, text="Click this link to complete the subscription: " + checkout_url)


                except (stripe.error.StripeError, mysql.connector.Error) as e:   # Handle Stripe and MySQL errors
                    print("An error occurred:", str(e))           
                    cur = connection.cursor()
                    cur.execute("SELECT remaining_duration FROM users WHERE id = %s", (user_id,))
                    result = cur.fetchone()
                    print(result)
                    cur.close() 
                    remaining_duration = result['remaining_duration']
                    bot.send_message(chat_id=update.effective_chat.id, text=f"You already have paid for the subscription.\nThe remaining duration is {remaining_duration} days.")
    ######################################################################################################################################################################
            else:
                # Check if the command is valid and get the response
                chat_id = message.chat.id
                if chat_id not in user_conversation or not user_conversation[chat_id]:
                    bot.reply_to(message, "Please start a conversation before using this command.")
                    return
                else:
                    # Use a dictionary to map commands to prompts
                    command_to_prompt = {
                    "/ask": "Please ask me a question.",
                    "/grammar": "Correct this conversation:" + user_conversation[message.chat.id][-1],
                    "/words": "What are the upper intermediate and advanced words in this conversation?:" + bot_conversation[message.chat.id][-1] + bot_conversation_q[message.chat.id][-1],
                    "/score": "Score this conversation like an IELTS marker:" + user_conversation[message.chat.id][-1], # Dorost shavad.
                    "/analyze": "Analyze this conversation like an IELTS marker:" + user_conversation[message.chat.id][-1]
                }
                    if message.text == "/text":
                        logger.info("Received /text command")
                        bot.reply_to(message, "Conversation History:\n" + conversation_history)
                    elif message.text == "/export":
                # Create a new Word document
                        doc = docx.Document()
                # Add a heading to the document
                        doc.add_heading("Conversation History", 0)
                # Loop through the conversation history and add each message to the document
                        doc.add_paragraph(conversation_history)
                # Save the Word document
                        doc.save(f"conversation_history.docx")
                        convert("conversation_history.docx")
                        with open("conversation_history.docx", "rb") as file:
                            bot.send_document(chat_id=message.chat.id, document=file)
                        with open("conversation_history.pdf", "rb") as file:
                            bot.send_document(chat_id=message.chat.id, document=file)
                        bot.reply_to(message, "Conversation history exported to Word & PDF.")
                # Delete the word file
                        os.remove("conversation_history.docx")
                        os.remove("conversation_history.pdf")
                    #elif message.text == "/x":
                    #    return
                    elif message.text in command_to_prompt:
                        prompt = command_to_prompt[message.text] 
                        get_response(prompt)
                    else:
                        bot.reply_to(message, "Sorry, I didn't understand that. Please enter /help for help options.")
                
        except Exception as e:        
                logger.exception("Error in handling: %s", str(e))
                print(f"Error in handling: {e}")
                bot.reply_to(message,"Oops! Something went wrong. Please try again later.")
    else:
        start(update, context)





























users = set()

# Create a function to handle incoming voice messages
@bot.message_handler(content_types=['voice'])
def handle_voice_message(update, context):
    global my_voice_path 
    user_id = update.effective_user.id
    chat_id = update.message.chat_id
    chat_member = bot.get_chat_member(chat_id=chat_id, user_id=user_id)
    if chat_member.status == 'member':
        message = update.message
        if message.from_user is None:
            return None

        user_id = message.from_user.id
        username = message.from_user.username
        users.add(user_id)

        # Get the file ID of the voice message
        file_id = message.voice.file_id

        # Get information about the voice message
        voice_info = bot.get_file(file_id)
        voice_url = f"https://api.telegram.org/file/bot{bot.token}/{voice_info.file_path}"

        # Download the voice message file
        response = requests.get(voice_url)

        with open('received_voice.ogg', 'wb') as voice_file:
            voice_file.write(response.content)

        # Convert the voice message to WAV format
        audio = AudioSegment.from_ogg('received_voice.ogg')
        audio.export('received_voice.wav', format='wav')





        ##################### process voice message  here ##############
        # Reply to the user
        # bot.reply_to(message, "Voice message received.")
        logger.info("Voice message received.")

        # Check if the user has an existing conversation
        if message.chat.id not in conversation:
            conversation[message.chat.id] = []
            user_conversation[message.chat.id] = []
            bot_conversation[message.chat.id] =[]
            bot_conversation_q[message.chat.id] =[]


        filename = "received_voice.wav"
        # Convert the speech to text using Google Speech Recognition
        r = sr.Recognizer()
        with sr.AudioFile(filename) as source:
            audio = r.record(source)  # Read the entire audio file
            try:
                confidence = r.recognize_google(audio, show_all=True)['alternative'][0]['confidence']
                if confidence > 0.75:
                    text_user = r.recognize_google(audio)
                    if "ielts" in context.user_data and  context.user_data["ielts"] == True:
                        if len(text_user) < 30:
                            bot.reply_to(message, "Your answer is too short. Please try again.")


                            return
                        else:
                            bot.reply_to(message, f"Text of your voice is:\n {text_user}.\n\n"
                                                f" Your pronounciation score is: {confidence:.2f}")   
                    else:
                        bot.reply_to(message, f"I heard: {text_user}.\n"
                                                f"Confidence: {confidence:.2f}")                                             
                else:
                    bot.reply_to(message,"Please come again. I couldn't recognize the audio.\n"
                                    "Keep your distance from noisy places and speak loudly and clearly.")
                    return 


            except sr.UnknownValueError:
                    bot.reply_to(message,"Please come again. I couldn't recognize the audio.\n"
                                    "Keep your distance from noisy places and speak loudly and clearly.")
                    return 
        # Add the user's message to the conversation history
        conversation[message.chat.id].append("User: "+text_user+"\n")
        user_conversation[message.chat.id].append("User: "+text_user)
        print("************************************************************************************")
        print("This username is chatting:", username)
        print("\nUser:",text_user)
######################################################### IELTS - PART ######################################################           
        if "ielts" in context.user_data and  context.user_data["ielts"] == True:            
            if "flag" in context.user_data and  context.user_data["flag"] == False:
                context.user_data["text"] = ""
                context.user_data["num_voices"] = 0
                context.user_data["flag"] = True
            
            print("Flag:",context.user_data["flag"])

            print("num_voices:", context.user_data["num_voices"])
            context.user_data["num_voices"]  += 1

            if context.user_data["num_voices"] >= 1:
                prompt_grammar_ielts = "Analyze these sentences gramatically and correct them.:\n" + text_user
                analyze_raw1 = openai.Completion.create(
                engine=model_id,
                prompt=prompt_grammar_ielts,
                max_tokens=80,
                n=1,
                stop=None,
                temperature=0.5,
                )
                analyze_answer1 = analyze_raw1.choices[0].text.strip()
                # Split the string by newline characters
                lines = analyze_answer1.split('\n')
                # Check if the first line has less than 20 characters
                if len(lines[0]) < 20:
                    # Remove the first two lines if the first line has less than 20 characters
                    lines = lines[2:]
                    # Join the remaining lines back together
                    analyze_answer1 = '\n'.join(lines)


                print("The corrected form of your answer is:\n" +analyze_answer1)
                bot.send_message(chat_id=user_id, text="We fixed the grammatical errors in your answer. The corrected form is:\n" + analyze_answer1)
                ######################################################################################################################
                # Wait for 60 seconds before making the next API call
                time.sleep(10)
                prompt_details_ielts = "Expand following sentences and add four lines with advanced words to make it professional:\n" + analyze_answer1
                analyze_raw2 = openai.Completion.create(
                engine=model_id,
                prompt=prompt_details_ielts,
                max_tokens=100,
                n=1,
                stop=None,
                temperature=0.6,
                )
                analyze_answer2 = analyze_raw2.choices[0].text.strip()
                print("The best answer for you:\n" +analyze_answer2)
                bot.send_message(chat_id=user_id, text="Make your answer better by adding these details:\n" + analyze_answer2)
                #####################################################################################################################
                #########For selected prompt 2
                if context.user_data["text"] == "Where are you from?": 
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "I was born and raised in a vibrant city located in [mention the country]."
                "However, my roots trace back to a culturally rich region within [mention specific area or community],"
                "which has significantly influenced my upbringing and values. Despite being from [city/country]"
                "I've had the opportunity to explore various cultures and lifestyles," 
                "which has broadened my perspectives and understanding of the world."
                " I consider myself fortunate to have experienced the diverse tapestry of my homeland,"
                " which has shaped me into the person I am today.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "I come from a city called [Name of City], which is situated in the [Name of Country]."
                "It's a vibrant and culturally diverse place nestled [mention geographical location if applicable, e.g.," 
                "along the coast/in the heart of the countryside]. The city is known for its rich history and modern developments,"
                "offering a blend of tradition and innovation. Growing up in [Name of City], I was exposed to various traditions,"
                "languages, and customs, which greatly influenced my worldview. The multicultural environment has shaped my appreciation"
                "for diversity and taught me the value of embracing different perspectives.Moreover, [Name of City] has a significant impact"
                "on my identity. Its bustling streets, historical landmarks, and the warmth of its people have all contributed to shaping who I am today."
                "Even though I've been fortunate to explore other places, the connection I feel with my hometown remains profound."
                "Overall, I'm proud to hail from such a culturally rich and dynamic place like [Name of City], "
                "as it has significantly shaped my outlook on life and instilled in me a deep sense of appreciation for diversity.")
                
                if context.user_data["text"]  == "Tell me a little bit about your hometown.":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "I come from a remarkable place called [Name of City/Town], nestled in the heart of [Name of Region/State/Province] in [Name of Country]. "
                "It's a city brimming with culture, history, and a unique charm that captivates anyone who visits. "
                "[Name of City/Town] is renowned for its [mention specific features such as landmarks, industries, natural beauty, or cultural aspects]. "
                "For instance, it boasts [highlight a prominent landmark, historical site, or natural attraction] that serves as a symbol of pride for the locals "
                "and a fascination for tourists. "
                "The atmosphere in [Name of City/Town] is incredibly vibrant, owing to its diverse population and a blend of modern amenities with traditional values. "
                "Its [mention any festivals, events, or traditions] are a testament to the rich heritage and the unity among its people. "
                "Growing up in this dynamic city has been a privilege. The community here is incredibly welcoming and supportive, fostering a sense of belonging "
                "and warmth that shaped my formative years. Additionally, the education system and the opportunities available have contributed significantly "
                "to my personal growth and aspirations. "
                "While I've had the chance to explore other places, the unique essence of [Name of City/Town] remains incomparable. Its influence on my values, "
                "perspectives, and sense of identity is profound. I'm truly proud to call [Name of City/Town] my hometown, a place that encapsulates the best "
                "of cultural richness and community spirit.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "I hail from a place called [Name of City/Town], nestled in the heart of [Name of Region/State/Province] in [Name of Country]. "
                "[Name of City/Town] is a melting pot of culture, history, and charm, drawing people in with its unique character. "
                "Its distinguishing features include [mention specific aspects like landmarks, industries, natural beauty, or cultural events]. "
                "For instance, [highlight a prominent landmark or a significant event] stands as a symbol of pride for locals and attracts visitors alike. "
                "The vibrant atmosphere of [Name of City/Town], stemming from its diverse population and a blend of modern conveniences with traditional values, "
                "creates an enchanting aura. Celebrations like [mention festivals or events] showcase our rich heritage and the unity within our community. "
                "Growing up in this dynamic city has been a privilege, where the community's warmth and support have profoundly influenced my formative years. "
                "The educational opportunities and resources available here have significantly contributed to my personal growth and aspirations. "
                "Despite having explored other places, the essence of [Name of City/Town] remains incomparable. Its influence on my values, perspectives, "
                "and sense of identity is profound. I take immense pride in calling [Name of City/Town] my hometown, a place that epitomizes cultural richness "
                "and a strong community spirit.")
                    
                if context.user_data["text"]  == "What do you like most about your hometown?":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "One of the aspects I admire most about my hometown, [Name of City/Town], is its profound sense of community. "
                "The camaraderie and unity among its residents create an environment that fosters warmth and support. "
                "The people here are incredibly welcoming, which contributes to a strong sense of belonging. "
                "Additionally, the rich cultural tapestry woven through various traditions and customs is something I deeply cherish. "
                "The diversity in [Name of City/Town] allows me to appreciate different perspectives and learn from a wide range of experiences. "
                "Moreover, the city's [mention any specific aspects like natural beauty, historical landmarks, or local cuisine] adds to its charm "
                "and makes it an enticing place to live in. The blend of [mention any unique characteristics, such as modernity with traditional values] "
                "makes it a dynamic and engaging place. Overall, the sense of community, cultural richness, and the unique blend of modernity with tradition "
                "are the standout features that make me genuinely fond of my hometown, [Name of City/Town].")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "What captivates me most about my hometown, [Name of City/Town], is its remarkable sense of community. "
                "The unparalleled warmth and inclusivity among the people here create an atmosphere of belonging that's truly exceptional. "
                "The city's ability to preserve its cultural heritage while embracing modernity is another aspect I deeply admire. "
                "Its [mention specific cultural aspects, landmarks, or traditions] stand as symbols of pride and heritage, fostering a deep sense of identity among the locals. "
                "Moreover, the diversity within [Name of City/Town] is enriching; it's a vibrant mosaic of different cultures, languages, and traditions. "
                "This diversity has expanded my horizons, teaching me the value of acceptance and respect for various perspectives. "
                "The support network within the community is incredible, and the opportunities available, both educationally and professionally, are commendable. "
                "Despite having explored other places, the genuine connections and the feeling of homecoming in [Name of City/Town] remain unmatched. "
                "Overall, the strong sense of community, rich cultural heritage, and the embracing nature of [Name of City/Town] are what I cherish most.")
                #####################################################################################################################
                #########For selected prompt 3
                if context.user_data["text"] ==  "Are you a student or do you work?" or "What do you study or what is your occupation?": 
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "I am currently enrolled as a student pursuing [mention your field or degree program] at [name of your educational institution]. Concurrently, I also engage in part-time work as [mention your job position or industry], which allows me to gain practical experience and apply the theoretical knowledge I acquire from my studies.\n\n"
                "Balancing academics with work has been a rewarding experience for me. As a student, I am exposed to diverse perspectives and academic challenges, while my work allows me to develop valuable skills in [mention skills related to your job] and understand the professional world better.\n\n"
                "I believe this dual commitment not only enhances my educational journey but also prepares me for the responsibilities of the workplace. It's a balancing act that I manage efficiently, ensuring that I excel in both realms while maintaining a healthy lifestyle and time for personal growth.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "I am presently employed as a [your job position] at [name of your workplace]. However, I am also actively engaged in furthering my education by taking [mention any courses or educational programs you're involved in]. This allows me to expand my knowledge and skills while balancing professional responsibilities.\n\n"
                "Working in [your field or industry] has provided me with practical insights and experiences that supplement my academic pursuits. Simultaneously, my dedication to learning enables me to stay updated with the latest advancements and trends in my field.\n\n"
                "I value this harmonious blend of work and education as it offers me a comprehensive understanding of my field, empowering me to contribute meaningfully in both academic and professional environments.")
                
                if context.user_data["text"] ==  "Why did you choose that field of study or work?": 
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "The choice of [mention your field of study/work] stemmed from my intrinsic fascination with [specific aspects of your field]. I have always been captivated by [mention what interests or motivates you in the field], which led me to pursue this path.\n\n"
                "I firmly believe that [your field of study/work] offers a platform where I can make a tangible impact. Its dynamic nature and potential to contribute positively to [mention any relevant societal or global issues] deeply resonate with my personal values and aspirations.\n\n"
                "Moreover, the prospect of continuous learning and innovation within [your field of study/work] excites me. The challenges presented by this field serve as opportunities for growth and creativity, allowing me to constantly evolve and adapt.\n\n"
                "I am driven by the passion to [mention your goals or what you aim to achieve] in this domain, and I am committed to leveraging my skills and knowledge to make meaningful contributions.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "The decision to pursue [mention your field of study/work] was influenced by a blend of personal passion and pragmatic considerations. I've harbored a profound interest in [specific aspects of your field] from a young age, which drove my curiosity to explore and understand its nuances.\n\n"
                "Moreover, I recognized the immense scope for growth and innovation in [your field of study/work]. I was drawn to its potential to address [mention any relevant challenges or issues], and I aspire to be part of the solution in tackling these challenges.\n\n"
                "I am also inspired by the diverse opportunities this field presents. Whether it's the chance to conduct groundbreaking research, effect positive societal change, or contribute innovative ideas, [mention your field of study/work] offers a platform where my skills align with my ambitions.\n\n"
                "Ultimately, my passion, coupled with the substantial impact I can make in this field, solidifies my commitment to excel and contribute meaningfully.")
                #####################################################################################################################
                #########For selected prompt 4
                if context.user_data["text"] ==  "What do you enjoy doing in your free time?": 
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "In my leisure time, I find solace in pursuing various activities that cater to both my creative inclinations and my quest for personal growth. Engaging in [mention a hobby or activity you enjoy] is one of my primary interests. It allows me to channel my creativity and express myself artistically.\n\n"
                "Furthermore, I am an avid reader, delving into a diverse range of literature spanning genres like [mention your preferred genres]. Reading not only broadens my knowledge but also nurtures my imagination and offers a fresh perspective on different cultures and ideas.\n\n"
                "I also cherish spending time outdoors, particularly in nature. Whether it's going for hikes, exploring parks, or simply taking a stroll, being amidst nature rejuvenates me and serves as a great way to unwind from the demands of daily life.\n\n"
                "Lastly, I am passionate about [mention any other activities or interests you have], which not only brings me joy but also allows me to continuously learn and develop new skills.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "My leisure moments are diverse, often revolving around activities that both relax and stimulate me. One of my main interests is [mention a hobby or activity you enjoy], where I find an outlet for my creativity and a sense of accomplishment in creating something tangible.\n\n"
                "Additionally, I am passionate about [mention another interest or hobby]. This pursuit not only provides me with a platform for physical exercise but also cultivates discipline and focus, contributing to a healthy mind and body.\n\n"
                "Exploration is another aspect I cherish during my free time. Whether it's discovering new cuisines, visiting art galleries, or immersing myself in cultural events, I find immense joy in broadening my horizons and gaining new experiences.\n\n"
                "Moreover, I allocate time for [mention any other interests or activities], as it serves as a constant source of learning and personal growth, ensuring that I am continuously evolving.")
                
                if context.user_data["text"] ==  "Do you have any hobbies or interests?": 
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "Yes, I have several hobbies and interests that I enthusiastically dedicate my time to. One of my foremost passions is [mention a hobby or interest]. I find great pleasure and fulfillment in [describe why you enjoy this hobby or interest, and its significance in your life].\n\n"
                "Moreover, I am deeply engrossed in [mention another hobby or interest]. Engaging in this activity not only serves as a means of relaxation but also challenges me to [explain how this interest positively impacts you, whether mentally, physically, or emotionally].\n\n"
                "Additionally, I have an affinity for [mention another hobby or interest], which allows me to [describe the benefits or personal growth gained from this hobby]. The diversity in my hobbies not only keeps me motivated and energized but also nurtures my continuous learning and development.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "Yes, I am passionate about various hobbies and interests that play a significant role in my life. To begin with, I thoroughly enjoy [mention a hobby or interest], which allows me to [explain how this hobby brings joy or adds value to your life]. It serves as a creative outlet and a way to unwind from daily routines.\n\n"
                "Additionally, I find great pleasure in [mention another hobby or interest]. Engaging in this activity not only fosters my curiosity but also helps me [describe the personal or intellectual growth it provides].\n\n"
                "Furthermore, I dedicate time to [mention another hobby or interest], which not only ignites my passion but also [explain how this hobby contributes to your personal development or well-being]. Overall, my diverse interests enrich my life, offering a balance between relaxation, self-improvement, and enjoyment.")
                
                if context.user_data["text"] ==  "How did you develop an interest in that particular hobby?": 
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                "My fascination with [mention the hobby or interest] began [mention when or how you started]. It all started when [explain the initial encounter or experience that sparked your interest in the hobby]. This moment ignited a curiosity within me that propelled me to delve deeper into the hobby.\n\n"
                "As I delved further, I discovered [mention specific aspects or details about the hobby that intrigued you]. The more I learned, the more engrossed I became, and I realized that this hobby resonated with my [mention any personal interests, values, or skills].\n\n"
                "I actively sought resources and guidance to enhance my skills in [mention the hobby], whether it was through [explain how you learned or improved in this hobby]. The process of learning and improvement has been immensely rewarding, solidifying my passion for this hobby.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                "My passion for [mention the hobby or interest] evolved organically over time. It all started when [describe the initial exposure or experience that triggered your curiosity in this hobby]. This experience left a lasting impression on me, igniting a desire to explore and learn more about [mention specific aspects of the hobby].\n\n"
                "As I delved deeper into this hobby, I found myself captivated by [mention particular elements or skills associated with the hobby]. The sense of fulfillment I derived from honing these skills and gaining knowledge motivated me to further invest my time and effort into mastering [mention the hobby].\n\n"
                "Moreover, I actively sought opportunities to engage with [mention any communities, mentors, or resources related to the hobby]. Interacting with like-minded individuals and learning from experienced practitioners has been instrumental in nurturing my interest and advancing my proficiency in this hobby.")
                #####################################################################################################################
                #########For selected prompt 5
                if context.user_data["text"] == "Do you have a large or small family?":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                                    "I come from a relatively small family. It consists of my parents, my younger sister, and myself. While we might be a small unit in terms of immediate family members, our bond is incredibly strong and tight-knit.\n\n"
                                    "Although we don't have an extensive extended family, our smaller size has its advantages. It allows for more intimate connections and closer relationships among us. We value quality time together and prioritize supporting each other through every phase of life.\n\n"
                                    "While we might not have a large number of relatives, the love, understanding, and encouragement within our small family unit create a warm and nurturing environment that I cherish deeply.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                                    "I belong to a relatively large family, comprising my parents, two siblings - an older brother and a younger sister - and my paternal grandparents. Our family gatherings are always vibrant and lively, filled with laughter and shared stories.\n\n"
                                    "Having a larger family means there are more individuals to learn from and share experiences with. I have the privilege of receiving guidance not only from my parents but also from my grandparents, whose wisdom and experiences are invaluable to me.\n\n"
                                    "While managing a larger family can be challenging at times, the support and sense of belonging within our extensive network of relatives are unparalleled. We stand by each other during both joyous celebrations and challenging times, fostering a strong sense of unity and togetherness.")

                if context.user_data["text"] == "Are you close to your family members?":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                                    "Yes, I am very close to my family members. We share a strong and tight bond that has been nurtured over the years through mutual respect, support, and understanding.\n\n"
                                    "I consider my family to be my pillars of strength. We communicate openly, share our joys, and support each other through challenges. This closeness has enabled us to build trust and a sense of unity within our family unit.\n\n"
                                    "Furthermore, our close-knit relationships extend beyond just immediate family. We often gather for family events, celebrations, and even during routine occasions, which further strengthens our connection and reinforces the importance of family ties in our lives.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                                    "I have a very close relationship with my family members, which I cherish immensely. We have cultivated a strong bond founded on mutual respect, love, and unwavering support for each other.\n\n"
                                    "We maintain regular communication despite our individual commitments, and this constant connection plays a crucial role in preserving our closeness. Sharing both joyful moments and supporting one another during challenging times has solidified our familial ties.\n\n"
                                    "Moreover, our family traditions and gatherings are essential in reinforcing our bond. Whether it's celebrating special occasions together or simply spending quality time, these moments foster a sense of togetherness and strengthen our relationships.")

                if context.user_data["text"] == "What activities do you enjoy doing with your family?":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                                    "My family and I engage in various activities that strengthen our bond and create lasting memories. One of our favorite pastimes is having regular family dinners. Gathering around the table allows us to catch up on each other's lives, share anecdotes, and simply enjoy quality time together.\n\n"
                                    "We also enjoy outdoor activities like hiking or picnics in nature. Exploring the outdoors not only brings us closer but also allows us to appreciate the beauty of nature while engaging in physical activities.\n\n"
                                    "Additionally, playing board games or engaging in friendly competitions such as charades or card games brings out laughter and creates a joyful atmosphere within our family. These activities encourage teamwork and friendly rivalry, making our family time fun and lively.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                                    "Our family revels in a variety of activities that foster our unity and create joyful moments. One of our favorite shared hobbies is cooking together. We often gather in the kitchen, preparing family recipes, trying new dishes, and bonding over the art of cooking.\n\n"
                                    "Furthermore, we have a shared passion for outdoor adventures. Whether it's cycling through scenic routes, going on nature walks, or planning day trips, these activities allow us to connect with nature and each other in a refreshing environment.\n\n"
                                    "Moreover, volunteering for community service initiatives is something we deeply value. Engaging in philanthropic activities together not only strengthens our family bond but also instills in us a sense of gratitude and empathy for others.")
                #####################################################################################################################
                #########For selected prompt 6
                if context.user_data["text"] == "What does a typical day in your life look like?":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                                    "My typical day revolves around a balance of personal, academic/work, and leisure activities. I usually start my day early, engaging in a morning routine consisting of meditation and exercise to invigorate both my mind and body.\n\n"
                                    "Following this, I dedicate a few hours to my academic/work commitments. Whether attending classes, working on projects, or handling professional tasks, I allocate focused time to ensure productivity and progress.\n\n"
                                    "During breaks, I often engage in hobbies such as reading or pursuing artistic endeavors, which provide mental relaxation and a creative outlet.\n\n"
                                    "In the evenings, I prioritize spending time with my family, enjoying dinner together and sharing our daily experiences. Afterward, I allocate some time for personal relaxation, be it watching a movie, listening to music, or unwinding with a book before bedtime.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                                    "A typical day in my life is a harmonious blend of responsibilities, learning, and personal rejuvenation. I kickstart my mornings with a brisk workout session, which energizes me for the day ahead.\n\n"
                                    "Following this, I dedicate a significant portion of my day to academic pursuits or professional responsibilities. This involves attending classes, engaging in research, or handling work-related tasks, ensuring that I make progress towards my goals.\n\n"
                                    "During breaks or free moments, I immerse myself in hobbies like painting or playing a musical instrument. These creative intervals act as refreshing pauses, allowing me to recharge and stimulate my creativity.\n\n"
                                    "Evenings are reserved for quality time with family or friends. We often gather for dinner, engaging in meaningful conversations and sharing experiences, which fosters a sense of belonging and connection.")
                
                if context.user_data["text"] == "Do you have any morning or evening routines?":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                                    "Yes, I have established specific morning and evening routines that help me maintain a sense of structure and productivity in my daily life.\n\n"
                                    "In the mornings, I begin by setting a positive tone for the day. I wake up early to practice mindfulness through meditation or yoga, which allows me to center my thoughts and prepare for the day ahead. Following this, I engage in a light workout session to boost my energy levels.\n\n"
                                    "In the evenings, I focus on winding down and relaxation. I usually allocate time for reading or journaling, which helps me reflect on the day's events and unwind mentally. Additionally, I prioritize disconnecting from screens to ensure a restful night's sleep.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                                    "Yes, I firmly believe in the importance of establishing morning and evening routines for a balanced and productive lifestyle.\n\n"
                                    "In the mornings, I kickstart my day with a focused approach. I begin by dedicating time to a mindful practice, such as meditation or a brief reflection, setting positive intentions for the day. This is followed by a nutritious breakfast and a short exercise routine, which energizes me for the day's tasks.\n\n"
                                    "Conversely, my evening routine is geared towards winding down and relaxation. I allocate time for leisure activities, such as reading a book or listening to calming music, to ease into a more tranquil state of mind. Additionally, I ensure to disconnect from electronic devices an hour before bedtime to promote better sleep quality.")
                
                if context.user_data["text"] == "How do you usually spend your weekends?":
                    bot.send_message(chat_id=user_id, text=f"This is the first answer with the highest score:\n"
                                    "I consider weekends as a time for rejuvenation and a break from the routine. On Saturdays, I often engage in outdoor activities to make the most of the day. I might go hiking or take leisurely walks in parks, appreciating nature's beauty and getting some fresh air.\n\n"
                                    "Conversely, Sundays are more relaxed and focused on personal pursuits. I allocate time for hobbies that I might not get to enjoy during the week. Whether it's painting, playing a musical instrument, or engaging in creative writing, these activities provide a sense of fulfillment and relaxation.\n\n"
                                    "Additionally, weekends are also about spending quality time with family and friends. I often plan gatherings or outings with loved ones, which fosters deeper connections and creates cherished memories.")
                    bot.send_message(chat_id=user_id, text=f"This is the second answer with the highest score:\n"
                                    "My weekends are a blend of relaxation, productivity, and social engagements. On Saturdays, I often devote time to personal development. This might involve pursuing online courses, reading informative literature, or engaging in activities that stimulate my mind.\n\n"
                                    "Conversely, Sundays are reserved for leisure and social connections. I prioritize spending quality time with family and friends, whether it's sharing a meal, exploring new places, or simply enjoying engaging conversations.\n\n"
                                    "Moreover, I appreciate incorporating some self-care routines into my weekends, such as practicing yoga or going for a spa treatment. These activities help me unwind and recharge for the upcoming week.")


            conversation_IELTS1 = "\n".join(conversation[message.chat.id])
            
            user_id = update.effective_user.id
            chat_id = update.effective_chat.id   
            if context.user_data["num_voices"] == 1:
                context.user_data["text"] = function2(user_id, chat_id)     
            if context.user_data["num_voices"] == 2:
                context.user_data["text"] = function3(user_id, chat_id)
            if context.user_data["num_voices"] == 3:
                context.user_data["text"] = function4(user_id, chat_id)
            if context.user_data["num_voices"] == 4:
                context.user_data["text"] = function5(user_id, chat_id)
            if context.user_data["num_voices"] == 5:
                conversation_IELTS1, context.user_data["text"] = function6(user_id, chat_id)
            if context.user_data["num_voices"] == 6:
                marker(user_id, chat_id, conversation_IELTS1)
            
            #context.user_data["ielts"] = False
            
        
    ##################################################################################################################################
        else:
            special_hello_words = ["how are you"]
            special_introduce_words =["what is your name", "what's your name" ]
            special_curse_words = [ "shit", "bitch", "motherfucker", "motherfuker", "cunt",
                                    "asshole", "fuck", "piss off", "dick head" ,
                                    "son of a bitch", "bastard", "damn", "shut up"
                                    ]

            # Check for specific words or phrases
            if any(word in text_user.lower() for word in special_hello_words):
                response_text_s = "Hello! I am fine. Thank you ,and you? How can I assist you?"
            elif any(word in text_user.lower() for word in special_introduce_words):
                response_text_s = " My name is Elisa, and you?"
            elif any(word in text_user.lower() for word in special_curse_words):
                #response_text_s ="It is non of your business, you're such a bitch! Fuck you!!"

                response_curse_list = [ "If you have a problem with me, write it on a piece of paper, fold it and stick it up your ass.",
                                        "You are proof that evolution CAN go in reverse.",
                                        "I love what you've done with your hair. How do you get it to come out of the nostrils like that?",
                                        "Bastard, Do you still love nature, despite what it did to you?",
                                        "Just because you have one doesn't mean you have to act like one.",
                                        "Which sex position produces the ugliest children? Ask your mother.",
                                        "Maybe you should eat make up so you can be pretty on the inside, bitch.",
                                        "Shut up, you'll never be the man your mother is.",
                                        "Is your ass jealous of the amount of shit that just came out of your mouth?",
                                        "You're so ugly; when your mom dropped you off at school she got a fine for littering.",
                                        "There's only one problem with your face, I can see it.",
                                        "If you were twice as smart, you'd still be stupid.",
                                        "Bitch, How did you get here? Did someone leave your cage open?",
                                        "Asshole, I would slap you but that will be animal abuse.",
                                        "If I had a penny for every brain cell you have, I'd have nothing.",
                                        "I called your boyfriend gay and he hit me with his purse.",
                                        "From the moment I first saw you, I knew I wanted to spend the rest of my life avoiding you.",
                                        "If I wanted to kill myself, I would climb up your ego and jump down to your IQ level.",
                                        "Motherfucker! remember when I asked for your opinion? Me neither.",
                                        "It looks like your face caught on fire and someone tried to put it out with a hammer.",
                                        "I'm glad to see you're not letting your education get in the way of your ignorance.",
                                        "The last time I saw something like you, I flushed it.",
                                        "You're like the universe, constantly expanding.",
                                        "Save your breath... You'll need it to blow up your date.",
                                        "Bastard, You have the perfect face for radio.",
                                        "It is non of your business, you're such a bitch! Fuck you!!"
                                        "Bitch, you are like Monday, nobody likes you."
                                        ]
                # Randomly select a prompt
                selected_response = random.choice(response_curse_list)
                response_text_s = selected_response




            else:
                # Call the OpenAI API to generate a response ##########################################################Do not complete or repeat this message. it then tell me your attractive idea about this message.
                #Answer to this question: //failed
                #Tell me another idea about this message: //failed
                # Answer or elaborate on my comment:
                # Tell me your attractive opinion 
                prompt = "Answer my comment:" + text_user
                response = openai.Completion.create(
                engine=model_id,
                prompt=prompt,
                max_tokens=80,
                n=1,
                stop=None,
                temperature=0.5,
                )
                # Get the response text from the OpenAI API
                response_text = response.choices[0].text.strip()
                lines = response_text.split("\n")
                if len(lines) >= 2:
                    if lines[0].strip() == "" or lines[1].strip() == "":
                        response_text = "".join(lines[2:])
                # Convert the response text to speech using gTTS
                # You can change lang="en" for differet languages 
                tts = gTTS(text=response_text, lang="en", tld="com", slow=False)
                tts.save(os.path.join(my_voice_path, "custom_voice.mp3"))

                # Send the response text as an audio message
                with open(os.path.join(my_voice_path, "custom_voice.mp3"), 'rb') as voice:
                    bot.send_audio(message.chat.id, voice)

                # Delete the file after sending
                os.remove(os.path.join(my_voice_path, "custom_voice.mp3"))

                # Add the bot's response to the conversation history
                
                conversation[message.chat.id].append("Elisa: " + response_text)
                bot_conversation[message.chat.id].append("Elisa: "+ response_text)
                print("Elisa1:",response_text)
                # Send the conversation history as a text message
                conversation_history = "\n".join(conversation[message.chat.id])

                
                # Call the OpenAI API to generate a response ##########################################################Do not complete or repeat this message. it then tell me your attractive idea about this message.
                #Tell me another idea about this message  //failed 
                #Show me your disagreement with my comment
                #Tell me your attractive idea about this message //correct
                #Tell me your attractive opinion about my comment
                #Tell me your different perspective about my comment
                #Explain the reasoning behind my comment
                #Explain any potential drawbacks or challenges to my comment
                
                prompts = [
                "Tell me your attractive opinion about my comment:",
                "Tell me your different perspective about my comment:",
                "Explain the reasoning behind my comment:"
                #"Explain any potential drawbacks or challenges to my comment:"
                ]

                # Randomly select a prompt
                selected_prompt = random.choice(prompts)
                prompt = selected_prompt + text_user
                #print("Prompt is:   ", prompt)
                response = openai.Completion.create(
                engine=model_id,
                prompt=prompt,
                max_tokens=80,
                n=1,
                stop=None,
                temperature=0.5,
                )
                # Get the response text from the OpenAI API
                response_text = response.choices[0].text.strip()
                lines = response_text.split("\n")
                if len(lines) >= 2:
                    if lines[0].strip() == "" or lines[1].strip() == "":
                        response_text = "".join(lines[2:])
                # Convert the response text to speech using gTTS
                # You can change lang="en" for differet languages 
                tts = gTTS(text=response_text, lang="en", tld="com", slow=False)
                tts.save(os.path.join(my_voice_path, "custom_voice.mp3"))

                # Send the response text as an audio message
                with open(os.path.join(my_voice_path, "custom_voice.mp3"), 'rb') as voice:
                    bot.send_audio(message.chat.id, voice)

                # Delete the file after sending
                os.remove(os.path.join(my_voice_path, "custom_voice.mp3"))

                # Add the bot's response to the conversation history
                
                conversation[message.chat.id].append("Elisa: " + response_text)
                bot_conversation[message.chat.id].append("Elisa: "+ response_text)
                print("Elisa1:",response_text)
                # Send the conversation history as a text message
                conversation_history = "\n".join(conversation[message.chat.id])




                ################################################################################################Do not complete or repeat this message. 
                #Ask me a new attractive related question in the following:
                #Ask me another attractive question related to this text: //failed
                #Ask me another question related to this text: //failed
                #Ask me another question related to my comment:
                prompt_q = "Ask me question related to:" + text_user
                response_q = openai.Completion.create(
                engine=model_id,
                prompt=prompt_q,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
                )
                # Get the response text from the OpenAI API
                response_text1 = response_q.choices[0].text.strip()
                lines = response_text1.split("\n")
                if len(lines) >= 2:
                    if lines[0].strip() == "" or lines[1].strip() == "":
                        response_text1 = "".join(lines[2:])
                
                #toggled = False
                #if toggled == False:
                #    response_text1 = "\n".join(response_text1.split("\n")[2:])
                #    toggled = True
                
                # Convert the response text to speech using gTTS
                # You can change lang="en" for differet languages 
                tts1 = gTTS(text=response_text1, lang="en", tld="com", slow=False)
                tts1.save(os.path.join(my_voice_path, "custom_voice1.mp3"))

                # Send the response text as an audio message
                with open(os.path.join(my_voice_path, "custom_voice1.mp3"), 'rb') as voice:
                    bot.send_audio(message.chat.id, voice)
                
                # Delete the file after sending
                os.remove(os.path.join(my_voice_path, "custom_voice1.mp3"))

                # Add the bot's response to the conversation history
                
                conversation[message.chat.id].append("Elisa: " + response_text1)
                bot_conversation_q[message.chat.id].append("Elisa: "+ response_text1)
                print("Elisa2:",response_text1)
                # Send the conversation history as a text message
                conversation_history = "\n".join(conversation[message.chat.id])
                return
            ##########################################################################################################################
            # Convert the response text to speech using gTTS
            # You can change lang="en" for differet languages 
            tts = gTTS(text=response_text_s, lang="en", tld="com", slow=False)
            tts.save(os.path.join(my_voice_path, "S_voice.mp3"))

            # Send the response text as an audio message
            with open(os.path.join(my_voice_path, "S_voice.mp3"), 'rb') as voice:
                bot.send_audio(message.chat.id, voice)

            # Delete the file after sending
            os.remove(os.path.join(my_voice_path, "S_voice.mp3"))

            # Add the bot's response to the conversation history
            
            conversation[message.chat.id].append("Elisa: " + response_text_s)
            bot_conversation[message.chat.id].append("Elisa: "+ response_text_s)
            print("Elisa1:",response_text_s)
            # Send the conversation history as a text message
            conversation_history = "\n".join(conversation[message.chat.id])

    else:
        start(update, context)

    # We can use these two lines for showing the conversation history
    # bot.reply_to(message, "The last Conversation:\n" + conversation[message.chat.id][-1])
    # bot.reply_to(message, "Conversation History:\n" + conversation_history)














def function1(user_id, chat_id):
    part1_1 = [ "Can you tell me your full name, please?\n",
                "What should I call you?"
    ]
    # Randomly select a prompt
    selected_prompt1 = random.choice(part1_1)
    response_text3 = selected_prompt1
    print(response_text3)   
    
    bot.send_message(chat_id=user_id, text=f"IELTS exam simulator:\n"
                    "Speaking part 1 \n"
                    "an introduction and interview\n"
                    "You have 5 minutes to answer 6 questions.\n")                
    bot.send_message(chat_id=user_id, text=f"Question 1:")
    # Convert the response text to speech using gTTS
    # You can change lang="en" for different languages 
    tts2 = gTTS(text=response_text3, lang="en", tld="com", slow=False)
    tts2.save(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    # Send the response text as an audio message
    with open(os.path.join(my_voice_path, "IELTS_voice.mp3"), 'rb') as voice:
        bot.send_audio(chat_id, voice)

    # Delete the file after sending
    os.remove(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    conversation[chat_id].append("Interviewer: " + response_text3)                
    conversation_IELTS1 = "\n".join(conversation[chat_id])




def function2(user_id, chat_id):
    bot.send_message(chat_id=user_id, text=f"The answer with the highest score is:\n"
                                            "My full name is [Your Full Name].\n"
                                            "You can call me [Preferred Name or Nickname].")
    part1_2 = [ "Where are you from?",
            "Tell me a little bit about your hometown.",
            "What do you like most about your hometown?"
    ]
    # Randomly select a prompt
    selected_prompt2 = random.choice(part1_2)
    response_text3 = selected_prompt2 
    print(response_text3)               
    bot.send_message(chat_id=user_id, text=f"Question 2:")
    # Convert the response text to speech using gTTS
    # You can change lang="en" for different languages 
    tts2 = gTTS(text=response_text3, lang="en", tld="com", slow=False)
    tts2.save(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    # Send the response text as an audio message
    with open(os.path.join(my_voice_path, "IELTS_voice.mp3"), 'rb') as voice:
        bot.send_audio(chat_id, voice)

    # Delete the file after sending
    os.remove(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    conversation[chat_id].append("Interviewer: " + response_text3)                
    conversation_IELTS1 = "\n".join(conversation[chat_id])
    return selected_prompt2




def function3(user_id, chat_id):
    part1_3 = [  "Are you a student or do you work?",
                "What do you study or what is your occupation?",
                "Why did you choose that field of study or work?"
    ]
    # Randomly select a prompt
    selected_prompt3 = random.choice(part1_3)
    response_text3 = selected_prompt3 
    print(response_text3)               
    bot.send_message(chat_id=user_id, text=f"Question 3:")
    # Convert the response text to speech using gTTS
    # You can change lang="en" for different languages 
    tts2 = gTTS(text=response_text3, lang="en", tld="com", slow=False)
    tts2.save(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    # Send the response text as an audio message
    with open(os.path.join(my_voice_path, "IELTS_voice.mp3"), 'rb') as voice:
        bot.send_audio(chat_id, voice)

    # Delete the file after sending
    os.remove(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    conversation[chat_id].append("Interviewer: " + response_text3)                
    conversation_IELTS1 = "\n".join(conversation[chat_id])
    return selected_prompt3





def function4(user_id, chat_id):
    part1_4 = [ "What do you enjoy doing in your free time?",
                "Do you have any hobbies or interests?",
                "How did you develop an interest in that particular hobby?"
    ]
    # Randomly select a prompt
    selected_prompt4 = random.choice(part1_4)
    response_text3 = selected_prompt4 
    print(response_text3)               
    bot.send_message(chat_id=user_id, text=f"Question 4:")
    # Convert the response text to speech using gTTS
    # You can change lang="en" for different languages 
    tts2 = gTTS(text=response_text3, lang="en", tld="com", slow=False)
    tts2.save(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    # Send the response text as an audio message
    with open(os.path.join(my_voice_path, "IELTS_voice.mp3"), 'rb') as voice:
        bot.send_audio(chat_id, voice)

    # Delete the file after sending
    os.remove(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    conversation[chat_id].append("Interviewer: " + response_text3)                
    conversation_IELTS1 = "\n".join(conversation[chat_id])
    return selected_prompt4



def function5(user_id, chat_id):
    part1_5 = [ "Do you have a large or small family?",
                "Are you close to your family members?",
                "What activities do you enjoy doing with your family?"
    ]
    # Randomly select a prompt
    selected_prompt5 = random.choice(part1_5)
    response_text3 = selected_prompt5
    print(response_text3)               
    bot.send_message(chat_id=user_id, text=f"Question 5:")
    # Convert the response text to speech using gTTS
    # You can change lang="en" for different languages 
    tts2 = gTTS(text=response_text3, lang="en", tld="com", slow=False)
    tts2.save(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    # Send the response text as an audio message
    with open(os.path.join(my_voice_path, "IELTS_voice.mp3"), 'rb') as voice:
        bot.send_audio(chat_id, voice)

    # Delete the file after sending
    os.remove(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    conversation[chat_id].append("Interviewer: " + response_text3)                
    conversation_IELTS1 = "\n".join(conversation[chat_id])
    return selected_prompt5


def function6(user_id, chat_id):
    part1_6 = [ "What does a typical day in your life look like?\n",
                "Do you have any morning or evening routines?\n",
                "How do you usually spend your weekends?"
    ]
    # Randomly select a prompt
    selected_prompt6 = random.choice(part1_6)
    response_text3 = selected_prompt6 
    print(response_text3)               
    bot.send_message(chat_id=user_id, text=f"Question 6:")
    # Convert the response text to speech using gTTS
    # You can change lang="en" for different languages 
    tts2 = gTTS(text=response_text3, lang="en", tld="com", slow=False)
    tts2.save(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    # Send the response text as an audio message
    with open(os.path.join(my_voice_path, "IELTS_voice.mp3"), 'rb') as voice:
        bot.send_audio(chat_id, voice)

    # Delete the file after sending
    os.remove(os.path.join(my_voice_path, "IELTS_voice.mp3"))
    conversation[chat_id].append("Interviewer: " + response_text3)                
    conversation_IELTS1 = "\n".join(conversation[chat_id])
    return conversation_IELTS1, selected_prompt6
    





def marker(user_id, chat_id, conversation_IELTS1):
    #prompt_ielts1 ="Please analyze my performance in the speaking part of the IELTS exam based on the provided audio recording:\n" + conversation_IELTS1
    prompt_ielts1 = "Analyze my performance in the speaking part of the IELTS exam:\n" + conversation_IELTS1
    response_marker_raw1 = openai.Completion.create(
    engine=model_id,
    prompt=prompt_ielts1,
    max_tokens=80,
    n=1,
    stop=None,
    temperature=0.5,
    )
    response_marker1 = response_marker_raw1.choices[0].text.strip()
    print("response_marker1: " + response_marker1)

    # Wait for 60 seconds before making the next API call
    time.sleep(30)




    prompt_ielts2 = "Provide scores tables of grammar,vocabulary,fluency,coherence in the speaking part of the IELTS exam:\n" + conversation_IELTS1 
    
    response_marker_raw2 = openai.Completion.create(
    engine=model_id,
    prompt=prompt_ielts2,
    max_tokens=80,
    n=1,
    stop=None,
    temperature=0.5,
    )
    response_marker2 = response_marker_raw2.choices[0].text.strip()
    print("response_marker2: " + response_marker2)

    # Wait for 60 seconds before making the next API call
    time.sleep(30)




    prompt_ielts3 = "Provide feedback on my weaknesses and areas of improvement in the speaking part of the IELTS exam:\n" + conversation_IELTS1
    response_marker_raw3 = openai.Completion.create(
    engine=model_id,
    prompt=prompt_ielts3,
    max_tokens=80,
    n=1,
    stop=None,
    temperature=0.5,
    )
    response_marker3 = response_marker_raw3.choices[0].text.strip()
    print("response_marker3: " + response_marker3)




    # Print response_text to debug
    response_marker_T = "\n"+"Marker analyze: " +  response_marker1 + "\n\nMarker scores: " + response_marker2 + "\n\nYour weaknesses: " + response_marker3 
    print(response_marker_T)
    # Convert the response text to speech using gTTS
    
    # You can change lang="en" for different languages 
    tts_marker = gTTS(text=response_marker_T , lang="en", tld="com", slow=False)
    tts_marker.save(os.path.join(my_voice_path, "ielts_marker.mp3"))
    # Send the response text as an audio message
    with open(os.path.join(my_voice_path, "ielts_marker.mp3"), 'rb') as voice:
        bot.send_audio(chat_id, voice)
    
    # Delete the file after sending
    os.remove(os.path.join(my_voice_path, "ielts_marker.mp3"))
    conversation[chat_id].append(response_marker_T)
    conversation_IELTS1 = "\n".join(conversation[chat_id])
    bot.send_message(chat_id=user_id, text=f"Speaking History:\n" + conversation_IELTS1)























# Define a rate-limited function using the "limits" decorator
@sleep_and_retry
@limits(calls=RATE_LIMIT, period=60)
def start(update: Update, context):
    logger.info("Received /start command")
    try:
        user = update.message.from_user
        if user is None:
            return None
        
        user_id = update.message.from_user.id
        users.add(user_id)
        username = update.message.from_user.username

        #total_users = len(users)
        print("************************************************************************************")
        chat_id = update.message.chat_id
        members_count = bot.get_chat_members_count(chat_id)
        print("Total number of members in the chat: ", members_count)
        #print("Total number of users: ", total_users)
        print("Welcome, username:", username)
        user_info = bot.get_chat_member(chat_id, user_id)
        firstname = user_info.user.first_name
        lastname = user_info.user.last_name
        if firstname is None:
            firstname = "N/A"
        if lastname is None:
            lastname = "N/A"

        print("User firstname & lastname: ", firstname + "   " + lastname)
        
        user_id = update.effective_user.id
        chat_member = bot.get_chat_member(chat_id=chat_id, user_id=user_id)
        if chat_member.status == 'member':
            update.message.reply_text("Hi,\n"
                                "Welcome to my AI-powered bot for learning English based on active conversation!\n"
                                "Please read /privacy policy carefully before using this Bot.\n" 
                                "By using the Bot, you agree to the terms of this Policy.\n" 
                                "If you do not agree with the terms of this Policy, do not use the Bot.\n"
                                #"Please /sign_in .\nIf you don't have a password, /sign_up first and then /sign_in.\n"
                                #"Just record your voice for conversation  just after /sign_in.\n"
                                "If you enter /help , you will see keys for different functions.\n"
                                "Just record your voice for conversation .\n"
                                "If you enter /start , you will return to this menu.\n\n"

                                "ZebraZebra owns all intellectual property rights, including copyright, in and to the content and design of this Bot, unless otherwise indicated. You may not reproduce, distribute, or otherwise use any of the content or design of this Bot without our prior written consent.\n"
                                "Any unauthorized use of the content or design of this Bot may be a violation of copyright, trademark, or other laws.\n"
                                "Copyright ©\n"
                                "year:2023 ,  Reza Ghasemi\n"
                                "All rights reserved.")
            return
        else:  # @zebra_test_431
            update.message.reply_text("Please join this channel:\n"
                                    "@ZebraZebraAI \n"
                                    #"@zebra_test_431 \n"
                                    "After joining this channel,\n"
                                    "enter the /start command again for using this bot.")
            return 

    except Exception as e:
        print(f"Error in start: {e}")
        logger.exception("Error in start: %s", str(e))
        update.message.reply_text("Oops! Something went wrong. Please try again later.")



#Start the Telegram bot
def main():
    try:
        updater = Updater(token=BOT_TOKEN, use_context=True)
        dispatcher = updater.dispatcher
        dispatcher.add_handler(CommandHandler('start', start))
        dispatcher.add_handler(MessageHandler(Filters.text, handle_message))
        dispatcher.add_handler(MessageHandler(Filters.voice, handle_voice_message))
        updater.start_polling()
        updater.idle()
    except Exception as e:
        logger.exception("Error in main: %s", str(e))


if __name__ == '__main__':
    main()
